from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt

from config import BOOK_SOURCE_URL, OUTPUT_DIR, REPORT_AUTHOR
from models import BookAnalysis


class ReportError(RuntimeError):
    pass


def _set_default_font(document: Document) -> None:
    style = document.styles["Normal"]
    style.font.name = "Aptos"
    style.font.size = Pt(11)


def create_word_report(
    analysis: BookAnalysis,
    destination: Path | None = None,
    report_author: str = REPORT_AUTHOR,
) -> Path:
    destination = destination or OUTPUT_DIR / "rapport_alice.docx"
    destination.parent.mkdir(parents=True, exist_ok=True)

    try:
        document = Document()
        _set_default_font(document)

        title = document.add_paragraph()
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = title.add_run(analysis.title)
        run.bold = True
        run.font.size = Pt(24)

        subtitle = document.add_paragraph()
        subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
        subtitle_run = subtitle.add_run("Analyse du premier chapitre")
        subtitle_run.italic = True
        subtitle_run.font.size = Pt(15)

        picture_paragraph = document.add_paragraph()
        picture_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        picture_run = picture_paragraph.add_run()
        picture_run.add_picture(str(analysis.processed_image_path), width=Inches(4.4))

        info = document.add_paragraph()
        info.alignment = WD_ALIGN_PARAGRAPH.CENTER
        info.add_run(f"Auteur du livre : {analysis.author}\n").bold = True
        info.add_run(f"Auteur du rapport : {report_author}").italic = True

        document.add_page_break()

        heading = document.add_heading("Distribution des longueurs des paragraphes", level=1)
        heading.runs[0].bold = True
        heading.runs[0].italic = True

        chart_paragraph = document.add_paragraph()
        chart_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        chart_paragraph.add_run().add_picture(str(analysis.chart_path), width=Inches(6.5))

        document.add_heading("Description et statistiques", level=2)
        stats = (
            f"Le premier chapitre contient {len(analysis.paragraph_word_counts)} paragraphes "
            f"et {analysis.total_words} mots. Le paragraphe le plus court contient "
            f"{analysis.min_words} mots et le plus long {analysis.max_words} mots. "
            f"La moyenne est de {analysis.average_words:.2f} mots par paragraphe. "
            "Les longueurs sont regroupées par dizaines inférieures, conformément à "
            "l'exemple du sujet. L'image de couverture a été recadrée, redimensionnée, "
            "puis complétée avec un logo noir et blanc pivoté."
        )
        document.add_paragraph(stats)

        source = document.add_paragraph()
        source.add_run("Source des données : ").bold = True
        source.add_run(f"Project Gutenberg - {BOOK_SOURCE_URL}\n")
        source.add_run("Mode d’accès au texte : ").bold = True
        source.add_run(f"{analysis.text_source}\n")
        source.add_run("Mode d’accès à l’image : ").bold = True
        source.add_run(analysis.image_source)

        document.save(destination)
        return destination
    except Exception as exc:
        raise ReportError(f"Création du document Word impossible : {exc}") from exc
